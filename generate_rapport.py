#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Génère le rapport Word (docx) du projet ENT EST Salé."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Couleurs palette ──────────────────────────────────────────────────────────
BLEU_FONCE  = RGBColor(0x0F, 0x29, 0x4D)   # #0F294D  navy
BLEU_MED    = RGBColor(0x1A, 0x56, 0xDB)   # #1A56DB  accent
BLEU_CLAIR  = RGBColor(0xE8, 0xF0, 0xFE)   # #E8F0FE  fond cellule
VERT        = RGBColor(0x05, 0x96, 0x69)   # #059669  succès
GRIS_TEXTE  = RGBColor(0x37, 0x41, 0x51)   # #374151
BLANC       = RGBColor(0xFF, 0xFF, 0xFF)
ORANGE      = RGBColor(0xD9, 0x77, 0x06)

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'none'))
        tag.set(qn('w:sz'), kwargs.get('sz', '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    if level == 1:
        run.font.color.rgb = BLEU_FONCE
        run.font.size = Pt(18)
        run.bold = True
    elif level == 2:
        run.font.color.rgb = BLEU_MED
        run.font.size = Pt(14)
        run.bold = True
    elif level == 3:
        run.font.color.rgb = GRIS_TEXTE
        run.font.size = Pt(12)
        run.bold = True
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after  = Pt(6)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = GRIS_TEXTE
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRIS_TEXTE
    p.paragraph_format.left_indent = Cm(level * 0.5 + 0.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_info_box(doc, title, content, color_hex='1A56DB'):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, color_hex.replace('#','') if color_hex.startswith('#') else color_hex)
    set_cell_border(cell, top='single', left='single', bottom='single', right='single', sz='6', color=color_hex.replace('#',''))
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.2)
    r = p.add_run(f"  {title}: ")
    r.bold = True
    r.font.color.rgb = BLANC
    r.font.size = Pt(10)
    r2 = p.add_run(content)
    r2.font.color.rgb = BLANC
    r2.font.size = Pt(10)
    doc.add_paragraph()
    return table

def add_service_table(doc, services):
    """Add a styled table for microservices."""
    tbl = doc.add_table(rows=1 + len(services), cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Microservice', 'Port', 'Rôles Autorisés', 'Fonctionnalité Principale']
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '0F294D')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = BLANC
        r.font.size = Pt(10)

    for row_i, (name, port, roles, desc) in enumerate(services):
        row = tbl.rows[row_i + 1]
        bg = 'E8F0FE' if row_i % 2 == 0 else 'FFFFFF'
        data = [name, port, roles, desc]
        for col_i, val in enumerate(data):
            cell = row.cells[col_i]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            r.font.color.rgb = GRIS_TEXTE
            if col_i == 0:
                r.bold = True
                r.font.color.rgb = BLEU_FONCE
    doc.add_paragraph()
    return tbl

def add_infra_table(doc, infra):
    tbl = doc.add_table(rows=1 + len(infra), cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Composant', 'Port(s)', 'Rôle']
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '059669')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = BLANC
        r.font.size = Pt(10)
    for row_i, (name, port, role) in enumerate(infra):
        row = tbl.rows[row_i + 1]
        bg = 'F0FDF4' if row_i % 2 == 0 else 'FFFFFF'
        for col_i, val in enumerate([name, port, role]):
            cell = row.cells[col_i]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(9.5)
            r.font.color.rgb = GRIS_TEXTE
            if col_i == 0:
                r.bold = True
    doc.add_paragraph()
    return tbl

def add_page_break(doc):
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
def build_document():
    doc = Document()

    # ── Marges ────────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.5)

    # ── COUVERTURE ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()

    uni = doc.add_paragraph()
    uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = uni.add_run("Université Mohammed V — Rabat")
    r.font.size = Pt(13); r.font.color.rgb = GRIS_TEXTE; r.bold = True

    ecole = doc.add_paragraph()
    ecole.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ecole.add_run("École Supérieure de Technologie de Salé (EST Salé)")
    r.font.size = Pt(13); r.font.color.rgb = BLEU_FONCE; r.bold = True

    filiere = doc.add_paragraph()
    filiere.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = filiere.add_run("Filière : Ingénierie des Applications Web et Mobile")
    r.font.size = Pt(11); r.font.color.rgb = GRIS_TEXTE

    module_p = doc.add_paragraph()
    module_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = module_p.add_run("Module : DevOps et Cloud")
    r.font.size = Pt(11); r.font.color.rgb = BLEU_MED; r.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Titre principal
    titre = doc.add_paragraph()
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titre.add_run("RAPPORT DE PROJET")
    r.font.size = Pt(28); r.font.color.rgb = BLEU_FONCE; r.bold = True

    sous_titre = doc.add_paragraph()
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sous_titre.add_run("ENT-EST-Salé")
    r.font.size = Pt(22); r.font.color.rgb = BLEU_MED; r.bold = True

    desc_titre = doc.add_paragraph()
    desc_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = desc_titre.add_run("Plateforme Numérique de Travail — Architecture Microservices")
    r.font.size = Pt(14); r.font.color.rgb = GRIS_TEXTE; r.italic = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Encadrant & Groupe
    info_tbl = doc.add_table(rows=2, cols=2)
    info_tbl.style = 'Table Grid'
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    labels = [("Encadrant :", "—"), ("Année universitaire :", "2025–2026")]
    for i, (lbl, val) in enumerate(labels):
        set_cell_bg(info_tbl.rows[i].cells[0], '0F294D')
        set_cell_bg(info_tbl.rows[i].cells[1], 'E8F0FE')
        p0 = info_tbl.rows[i].cells[0].paragraphs[0]
        r0 = p0.add_run(lbl); r0.bold = True; r0.font.color.rgb = BLANC; r0.font.size = Pt(11)
        p1 = info_tbl.rows[i].cells[1].paragraphs[0]
        r1 = p1.add_run(val); r1.font.color.rgb = GRIS_TEXTE; r1.font.size = Pt(11)

    doc.add_paragraph()

    membres_tbl = doc.add_table(rows=6, cols=2)
    membres_tbl.style = 'Table Grid'
    membres_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_cell_bg(membres_tbl.rows[0].cells[0], '0F294D')
    set_cell_bg(membres_tbl.rows[0].cells[1], '0F294D')
    membres_tbl.rows[0].cells[0].merge(membres_tbl.rows[0].cells[1])
    p = membres_tbl.rows[0].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MEMBRES DU GROUPE"); r.bold = True; r.font.color.rgb = BLANC; r.font.size = Pt(11)

    membres = ["Membre 1", "Membre 2", "Membre 3", "Membre 4", "Membre 5"]
    for i, m in enumerate(membres):
        row = membres_tbl.rows[i + 1]
        bg = 'E8F0FE' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
        p0 = row.cells[0].paragraphs[0]
        r0 = p0.add_run(f"Étudiant {i+1}"); r0.font.size = Pt(10); r0.font.color.rgb = GRIS_TEXTE
        p1 = row.cells[1].paragraphs[0]
        r1 = p1.add_run(m); r1.font.size = Pt(10); r1.font.color.rgb = BLEU_FONCE; r1.bold = True

    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = date_p.add_run("Avril 2026")
    r.font.size = Pt(11); r.font.color.rgb = GRIS_TEXTE; r.italic = True

    add_page_break(doc)

    # ── TABLE DES MATIÈRES (manuelle) ─────────────────────────────────────────
    add_heading(doc, "Table des Matières", level=1)
    toc = [
        ("1.", "Résumé Exécutif", "3"),
        ("2.", "Contexte et Objectifs du Projet", "3"),
        ("3.", "Architecture Générale", "4"),
        ("4.", "Infrastructure Technique", "6"),
        ("5.", "Microservices — Description Détaillée", "8"),
        ("6.", "API Gateway", "13"),
        ("7.", "Interface Utilisateur (Frontend)", "14"),
        ("8.", "Sécurité et Authentification", "16"),
        ("9.", "Observabilité — Prometheus & Grafana", "17"),
        ("10.", "Tests et Validation", "18"),
        ("11.", "Déploiement avec Docker Compose", "19"),
        ("12.", "Patterns et Conventions de Développement", "20"),
        ("13.", "Conclusion et Perspectives", "21"),
    ]
    tbl_toc = doc.add_table(rows=len(toc), cols=2)
    for i, (num, title, page) in enumerate(toc):
        bg = 'E8F0FE' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(tbl_toc.rows[i].cells[0], bg)
        set_cell_bg(tbl_toc.rows[i].cells[1], bg)
        p0 = tbl_toc.rows[i].cells[0].paragraphs[0]
        r = p0.add_run(f"  {num}  {title}"); r.font.size = Pt(10.5); r.font.color.rgb = GRIS_TEXTE
        p1 = tbl_toc.rows[i].cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p1.add_run(f"{page}  "); r2.font.size = Pt(10.5); r2.font.color.rgb = BLEU_MED; r2.bold = True

    add_page_break(doc)

    # ── 1. RÉSUMÉ EXÉCUTIF ────────────────────────────────────────────────────
    add_heading(doc, "1. Résumé Exécutif", level=1)
    add_para(doc, (
        "Le projet ENT-EST-Salé est une plateforme numérique de travail (Espace Numérique de Travail) "
        "développée dans le cadre du module DevOps et Cloud à l'École Supérieure de Technologie de Salé "
        "(Université Mohammed V, Rabat). Il s'agit d'une application web complète répondant aux besoins "
        "académiques de l'établissement : gestion des cours, des examens, du calendrier, des forums de discussion "
        "et des notifications."
    ))
    add_para(doc, (
        "Architecturée selon les principes des microservices, la plateforme repose sur dix services indépendants "
        "orchestrés via Docker Compose, exposant une API unifiée à travers un API Gateway sécurisé par JWT. "
        "L'identité et le contrôle d'accès sont délégués à Keycloak (OpenID Connect), tandis que la persistance "
        "est assurée par Apache Cassandra, la messagerie asynchrone par RabbitMQ et le stockage d'objets par MinIO. "
        "La stack d'observabilité Prometheus + Grafana offre une visibilité temps-réel sur l'ensemble des services."
    ))

    add_page_break(doc)

    # ── 2. CONTEXTE ET OBJECTIFS ──────────────────────────────────────────────
    add_heading(doc, "2. Contexte et Objectifs du Projet", level=1)

    add_heading(doc, "2.1 Contexte Académique", level=2)
    add_para(doc, (
        "Ce projet s'inscrit dans le cadre pédagogique du module DevOps et Cloud de la filière "
        "Ingénierie des Applications Web et Mobile. L'objectif est de mettre en pratique les compétences "
        "acquises en matière de conteneurisation, d'architecture distribuée, de CI/CD et d'observabilité, "
        "en construisant une application réelle et complète."
    ))

    add_heading(doc, "2.2 Problématique", level=2)
    add_para(doc, (
        "Les établissements d'enseignement supérieur manquent souvent d'une plateforme numérique unifiée "
        "regroupant la gestion des cours, des devoirs, des échanges académiques et du calendrier. "
        "L'ENT-EST-Salé répond à ce besoin en proposant une solution moderne, scalable et sécurisée."
    ))

    add_heading(doc, "2.3 Objectifs Techniques", level=2)
    objectifs = [
        "Concevoir une architecture microservices découplée et maintenable",
        "Implémenter une authentification et une autorisation robustes (Keycloak + JWT + RBAC)",
        "Assurer la persistance des données avec Apache Cassandra (base NoSQL distribuée)",
        "Gérer le stockage de fichiers avec MinIO (compatible S3)",
        "Mettre en place une messagerie asynchrone avec RabbitMQ pour les événements inter-services",
        "Développer une interface utilisateur React moderne avec design system shadcn/ui",
        "Déployer l'ensemble via Docker Compose avec healthchecks et dépendances déclaratives",
        "Instrumenter tous les services avec Prometheus et visualiser via Grafana",
        "Écrire des scripts de smoke-test pour valider l'ensemble du pipeline fonctionnel",
    ]
    for o in objectifs:
        add_bullet(doc, o)

    add_page_break(doc)

    # ── 3. ARCHITECTURE ───────────────────────────────────────────────────────
    add_heading(doc, "3. Architecture Générale", level=1)
    add_para(doc, (
        "La plateforme suit une architecture microservices à trois niveaux : le niveau présentation (frontend React), "
        "le niveau API (gateway + microservices Python FastAPI) et le niveau infrastructure (bases de données, "
        "message broker, stockage d'objets et observabilité). Tous les composants sont conteneurisés via Docker "
        "et interconnectés dans le réseau interne ent-network."
    ))

    add_heading(doc, "3.1 Vue d'ensemble des couches", level=2)

    layers = [
        ("COUCHE PRÉSENTATION", "0F294D", [
            "React 18 + TypeScript + Vite (build tool)",
            "Tailwind CSS + shadcn/ui (composants UI)",
            "React Router v6 (navigation SPA)",
            "Contexte d'authentification (auth-context)",
            "WebSocket natif pour le chat temps-réel",
        ]),
        ("COUCHE API / GATEWAY", "1A56DB", [
            "ms-api-gateway : point d'entrée unique (port 8008)",
            "Validation JWT déléguée à ms-auth-core",
            "CORS configuré pour les origines autorisées",
            "Retry automatique sur les GET (idempotents)",
            "Propagation des erreurs downstream avec codes HTTP corrects",
        ]),
        ("COUCHE MICROSERVICES", "059669", [
            "ms-auth-core (8010) — Validation JWKS Keycloak",
            "ms-identity-admin (8013) — Provisionnement utilisateurs",
            "ms-course-content (8011) — Gestion des cours (enseignants)",
            "ms-course-access (8012) — Accès aux cours (étudiants)",
            "ms-notification (8014) — Notifications SMTP + Cassandra",
            "ms-calendar-schedule (8015) — Calendrier académique",
            "ms-forum-chat (8016) — Forum + WebSocket chat",
            "ms-exam-assignment (8017) — Devoirs, soumissions, notes",
            "ms-ai-assistant (8018) — Assistant IA (Ollama/Llama3)",
        ]),
        ("COUCHE INFRASTRUCTURE", "D97706", [
            "Keycloak 26 — IAM / OpenID Connect (port 8080)",
            "PostgreSQL 16 — Backend Keycloak (port 5432)",
            "Apache Cassandra 5 — Base de données NoSQL (port 9042)",
            "RabbitMQ 4 — Message broker AMQP (ports 5672/15672)",
            "MinIO — Stockage objet S3-compatible (ports 9001/9002)",
            "Redis 7 — Cache distribué (port 6379)",
            "Mailpit — SMTP sandbox de développement (ports 1025/8025)",
            "Prometheus — Collecte de métriques (port 9090)",
            "Grafana — Tableaux de bord (port 3001)",
        ]),
    ]

    for layer_name, color, items in layers:
        tbl = doc.add_table(rows=1 + len(items), cols=1)
        tbl.style = 'Table Grid'
        set_cell_bg(tbl.rows[0].cells[0], color)
        p = tbl.rows[0].cells[0].paragraphs[0]
        r = p.add_run(f"  {layer_name}")
        r.bold = True; r.font.color.rgb = BLANC; r.font.size = Pt(11)
        for i, item in enumerate(items):
            bg = 'F8FAFC' if i % 2 == 0 else 'FFFFFF'
            set_cell_bg(tbl.rows[i+1].cells[0], bg)
            p2 = tbl.rows[i+1].cells[0].paragraphs[0]
            r2 = p2.add_run(f"    • {item}")
            r2.font.size = Pt(10); r2.font.color.rgb = GRIS_TEXTE
        doc.add_paragraph()

    add_heading(doc, "3.2 Flux de Communication", level=2)
    add_para(doc, (
        "Deux patterns de communication coexistent dans la plateforme :"
    ))
    add_bullet(doc, "Synchrone (REST/HTTP) : Le frontend appelle le gateway → le gateway valide le JWT auprès de ms-auth-core → le gateway proxifie la requête vers le microservice cible avec l'en-tête Authorization préservé.")
    add_bullet(doc, "Asynchrone (AMQP) : Les microservices publient des événements sur l'exchange RabbitMQ ent.events.topic. ms-notification consomme les événements user.*, course.*, asset.*, assignment.* et grade.* pour envoyer des emails et stocker les notifications en base.")

    add_heading(doc, "3.3 Enveloppe d'Événement RabbitMQ", level=2)
    add_para(doc, "Tous les événements publiés suivent un format standardisé garantissant la traçabilité :")
    code_tbl = doc.add_table(rows=1, cols=1)
    set_cell_bg(code_tbl.rows[0].cells[0], '1E293B')
    p = code_tbl.rows[0].cells[0].paragraphs[0]
    code = (
        '{\n'
        '  "event_id":       "<uuid>",\n'
        '  "event_type":     "calendar.event.v1",\n'
        '  "occurred_at":    "2026-04-24T10:00:00Z",\n'
        '  "producer":       "ms-calendar-schedule",\n'
        '  "correlation_id": "<uuid>",\n'
        '  "payload":        { ... }\n'
        '}'
    )
    r = p.add_run(code)
    r.font.name = 'Courier New'; r.font.size = Pt(9); r.font.color.rgb = BLANC
    doc.add_paragraph()

    add_page_break(doc)

    # ── 4. INFRASTRUCTURE ─────────────────────────────────────────────────────
    add_heading(doc, "4. Infrastructure Technique", level=1)
    add_para(doc, (
        "L'infrastructure de la plateforme est entièrement définie dans un fichier docker-compose.yml unique, "
        "composé de plus de 25 services avec dépendances, healthchecks et volumes persistants. "
        "Les variables d'environnement sont centralisées dans .env.compose."
    ))

    infra_data = [
        ("Keycloak 26.1.4",      "8080, 9000",  "Serveur d'identité OpenID Connect, gestion des realms, rôles et utilisateurs"),
        ("PostgreSQL 16",         "5432",        "Base relationnelle backend de Keycloak (données d'identité)"),
        ("Apache Cassandra 5.0",  "9042",        "Base NoSQL large-colonne pour tous les microservices applicatifs"),
        ("RabbitMQ 4.1.3",        "5672, 15672", "Broker de messages AMQP, exchange topic ent.events.topic"),
        ("MinIO (latest)",        "9001, 9002",  "Stockage objet S3-compatible : cours, soumissions, logs"),
        ("Redis 7-alpine",        "6379",        "Cache distribué, sessions, rate-limiting"),
        ("Mailpit v1.18",         "1025, 8025",  "Serveur SMTP sandbox pour tester les emails sans envoi réel"),
        ("Prometheus",            "9090",        "Collecte métriques /metrics toutes les 15s sur les 9 services"),
        ("Grafana",               "3001",        "Visualisation dashboards, datasource Prometheus auto-provisionnée"),
        ("Ollama",                "11434",       "Runtime LLM local (profil ai, GPU Nvidia optionnel)"),
    ]
    add_infra_table(doc, infra_data)

    add_heading(doc, "4.1 Keycloak — Gestion des Identités", level=2)
    add_para(doc, (
        "Keycloak est déployé en mode développement avec un realm dédié ent-est configuré automatiquement "
        "via les scripts bootstrap-realm.sh et bootstrap-users.sh exécutés au démarrage. "
        "Le realm définit :"
    ))
    add_bullet(doc, "Trois rôles applicatifs : admin, teacher, student")
    add_bullet(doc, "Trois utilisateurs de développement : admin1, teacher1, student1")
    add_bullet(doc, "Un client OIDC ent-gateway avec audience ent-gateway, ent-frontend")
    add_bullet(doc, "Endpoint JWKS utilisé par ms-auth-core pour la validation des JWT")

    add_heading(doc, "4.2 Apache Cassandra — Persistance NoSQL", level=2)
    add_para(doc, (
        "Cassandra est choisi pour ses performances en écriture et sa scalabilité horizontale. "
        "Un keyspace unique ent_est avec SimpleStrategy (replication_factor=1) contient les tables de tous "
        "les microservices. Chaque service crée ses tables à l'initialisation (IF NOT EXISTS) via le pattern "
        "_ensure_db(). L'authentification utilise PlainTextAuthProvider."
    ))

    tables_cass = [
        "user_profiles — Profils utilisateurs (ms-identity-admin)",
        "courses — Métadonnées des cours (ms-course-content)",
        "course_assets — Assets/fichiers par cours (ms-course-content)",
        "notifications — Notifications par utilisateur (ms-notification)",
        "calendar_events — Événements du calendrier (ms-calendar-schedule)",
        "forum_threads — Fils de discussion (ms-forum-chat)",
        "forum_messages — Messages des fils (ms-forum-chat)",
        "forum_thread_counters — Compteurs (ms-forum-chat)",
        "assignments — Devoirs publiés (ms-exam-assignment)",
        "submissions — Soumissions étudiants (ms-exam-assignment)",
        "submissions_by_assignment — Index inversé (ms-exam-assignment)",
    ]
    for t in tables_cass:
        add_bullet(doc, t)

    add_heading(doc, "4.3 RabbitMQ — Messagerie Asynchrone", level=2)
    add_para(doc, (
        "RabbitMQ gère la communication événementielle entre services. Un exchange topic durable "
        "ent.events.topic reçoit tous les événements. La clé de routage correspond aux deux premiers "
        "segments du type d'événement (ex: calendar.event pour calendar.event.v1). "
        "ms-notification est le principal consommateur, abonné à la queue q.notification.user."
    ))

    add_heading(doc, "4.4 MinIO — Stockage Objet", level=2)
    add_para(doc, (
        "MinIO est le système de stockage objet compatible S3 utilisé pour stocker les fichiers des cours "
        "et les soumissions d'étudiants. Trois buckets sont créés automatiquement : ent-courses (contenus pédagogiques), "
        "ent-uploads (soumissions d'examens) et ent-logs. Les URLs de téléchargement sont des presigned URLs "
        "avec TTL configurable (défaut 180s) générées par ms-course-access."
    ))

    add_page_break(doc)

    # ── 5. MICROSERVICES ──────────────────────────────────────────────────────
    add_heading(doc, "5. Microservices — Description Détaillée", level=1)
    add_para(doc, (
        "Chaque microservice suit une structure de fichiers identique : app/__init__.py, app/runtime.py "
        "(métriques Prometheus, logging structuré, gestionnaires d'erreurs), app/main.py (logique métier), "
        "requirements.txt et Dockerfile. Cette uniformité facilite la maintenance et l'onboarding."
    ))

    services_tbl = [
        ("ms-auth-core",          "8010", "Interne",          "Validation JWT via JWKS Keycloak, endpoint /auth/me"),
        ("ms-api-gateway",        "8008", "Public",           "Point d'entrée unique, guard JWT, proxy vers services"),
        ("ms-identity-admin",     "8013", "admin",            "CRUD utilisateurs Keycloak + profils Cassandra"),
        ("ms-course-content",     "8011", "teacher, admin",   "Création/modification cours + upload assets MinIO"),
        ("ms-course-access",      "8012", "student",          "Listing cours + presigned download URLs"),
        ("ms-notification",       "8014", "Tous (consumer)",  "RabbitMQ consumer, SMTP Mailpit, stockage Cassandra"),
        ("ms-calendar-schedule",  "8015", "teacher/admin ✍︎", "CRUD événements calendrier, événements RabbitMQ"),
        ("ms-forum-chat",         "8016", "Tous",             "Forum REST + WebSocket chat temps-réel par salle"),
        ("ms-exam-assignment",    "8017", "teacher/student",  "Devoirs, soumissions MinIO, notation, événements"),
        ("ms-ai-assistant",       "8018", "Tous",             "Chat IA, résumé, FAQ via Ollama/Llama3"),
    ]
    add_service_table(doc, services_tbl)

    # ms-auth-core
    add_heading(doc, "5.1 ms-auth-core — Cœur d'Authentification", level=2)
    add_para(doc, "Responsabilité unique : valider les JWT Bearer émis par Keycloak. Le service récupère les clés publiques depuis l'endpoint JWKS de Keycloak et les met en cache localement.")
    add_bullet(doc, "GET /auth/health — Healthcheck (utilisé par Docker et le gateway)")
    add_bullet(doc, "GET /auth/me — Décode et valide le JWT, retourne les claims complètes (sub, preferred_username, realm_access.roles, etc.)")
    add_bullet(doc, "GET /metrics — Endpoint Prometheus (via runtime.py)")

    # ms-identity-admin
    add_heading(doc, "5.2 ms-identity-admin — Administration des Identités", level=2)
    add_para(doc, "Permet aux administrateurs de créer des utilisateurs (Keycloak + profil Cassandra) et de modifier leurs rôles.")
    add_bullet(doc, "POST /admin/users — Créer un utilisateur dans Keycloak et insérer le profil en Cassandra")
    add_bullet(doc, "GET /admin/users — Lister les utilisateurs (avec pagination et recherche)")
    add_bullet(doc, "GET /admin/users/{id} — Détail d'un utilisateur")
    add_bullet(doc, "PATCH /admin/users/{id}/roles — Modifier les rôles realm d'un utilisateur")
    add_bullet(doc, "Publie un événement user.created.v1 sur RabbitMQ après création")

    # ms-course-content
    add_heading(doc, "5.3 ms-course-content — Contenu des Cours", level=2)
    add_para(doc, "Service réservé aux enseignants et administrateurs pour gérer le catalogue de cours et leurs contenus pédagogiques.")
    add_bullet(doc, "GET/POST /courses — Lister/créer des cours (métadonnées en Cassandra)")
    add_bullet(doc, "PUT/DELETE /courses/{id} — Modifier/supprimer un cours")
    add_bullet(doc, "POST /courses/{id}/assets — Upload de fichier (PDF, vidéo, etc.) vers MinIO")
    add_bullet(doc, "DELETE /courses/{id}/assets/{asset_id} — Supprimer un asset")
    add_bullet(doc, "Publie des événements course.created.v1 et asset.uploaded.v1 sur RabbitMQ")

    # ms-course-access
    add_heading(doc, "5.4 ms-course-access — Accès Étudiant aux Cours", level=2)
    add_para(doc, "Service en lecture seule pour les étudiants, délègue la récupération des métadonnées à ms-course-content via jeton interne.")
    add_bullet(doc, "GET /courses — Liste tous les cours disponibles")
    add_bullet(doc, "GET /courses/{id} — Détail d'un cours")
    add_bullet(doc, "POST /courses/{id}/download-links — Génère une presigned URL MinIO (TTL 180s) pour télécharger un asset")

    # ms-notification
    add_heading(doc, "5.5 ms-notification — Notifications", level=2)
    add_para(doc, "Service hybride : consommateur RabbitMQ en arrière-plan et API REST pour lire les notifications.")
    add_bullet(doc, "Consumer AMQP : écoute tous les événements applicatifs, envoie un email via Mailpit et stocke la notification en Cassandra")
    add_bullet(doc, "GET /notifications/{user_id} — Récupère les notifications d'un utilisateur")
    add_bullet(doc, "PATCH /notifications/{user_id}/{id}/read — Marque comme lue")
    add_bullet(doc, "PATCH /notifications/{user_id}/read-all — Marque toutes comme lues")

    # ms-calendar-schedule
    add_heading(doc, "5.6 ms-calendar-schedule — Calendrier Académique", level=2)
    add_para(doc, "Gestion des événements du calendrier académique (cours, examens, réunions) avec filtrage par mois et par module.")
    add_bullet(doc, "POST /calendar/events — Créer un événement (teacher/admin)")
    add_bullet(doc, "GET /calendar/events — Lister avec filtres ?month=YYYY-MM et ?module_code=X")
    add_bullet(doc, "GET /calendar/events/{id} — Détail d'un événement")
    add_bullet(doc, "PATCH /calendar/events/{id} — Modifier un événement")
    add_bullet(doc, "DELETE /calendar/events/{id} — Supprimer un événement")
    add_bullet(doc, "Publie calendar.event.created.v1 sur RabbitMQ")

    # ms-forum-chat
    add_heading(doc, "5.7 ms-forum-chat — Forum et Chat Temps-Réel", level=2)
    add_para(doc, "Service dual-mode combinant une API REST pour les fils de discussion et un WebSocket pour le chat en direct par salle (module).")
    add_bullet(doc, "POST /forum/threads — Créer un fil de discussion")
    add_bullet(doc, "GET /forum/threads — Lister les fils (filtrable par module_code)")
    add_bullet(doc, "GET /forum/threads/{id} — Détail d'un fil avec messages")
    add_bullet(doc, "POST /forum/threads/{id}/messages — Poster un message REST dans un fil")
    add_bullet(doc, "WebSocket ws://host:8016/chat/ws?token=<jwt>&room=<module> — Chat temps-réel par salle")
    add_bullet(doc, "Compteurs Cassandra pour le nombre de messages par fil")

    # ms-exam-assignment
    add_heading(doc, "5.8 ms-exam-assignment — Devoirs et Examens", level=2)
    add_para(doc, "Gestion complète du cycle de vie des devoirs : création par l'enseignant, soumission par l'étudiant (fichier ou texte), correction et notation.")
    add_bullet(doc, "POST /assignments — Créer un devoir (teacher/admin)")
    add_bullet(doc, "GET /assignments — Lister les devoirs (filtres module_code, status)")
    add_bullet(doc, "POST /assignments/{id}/submissions — Soumettre un fichier (upload MinIO) ou texte")
    add_bullet(doc, "GET /assignments/{id}/submissions — Lister les soumissions")
    add_bullet(doc, "POST /assignments/{id}/submissions/{sid}/grade — Noter une soumission")
    add_bullet(doc, "Publie assignment.created.v1 et grade.assigned.v1 sur RabbitMQ")

    # ms-ai-assistant
    add_heading(doc, "5.9 ms-ai-assistant — Assistant IA", level=2)
    add_para(doc, "Service IA basé sur Ollama (runtime LLM local) avec le modèle Llama3. Accessible via profil Docker Compose 'ai' (nécessite GPU Nvidia optionnel).")
    add_bullet(doc, "POST /ai/chat — Chat conversationnel avec le modèle LLM")
    add_bullet(doc, "POST /ai/summarize — Résumé automatique d'un document ou texte")
    add_bullet(doc, "POST /ai/faq/generate — Génération de FAQ à partir d'un contenu")
    add_bullet(doc, "Rate limiting configurable (défaut 20 req/min)")

    add_page_break(doc)

    # ── 6. API GATEWAY ────────────────────────────────────────────────────────
    add_heading(doc, "6. API Gateway", level=1)
    add_para(doc, (
        "Le ms-api-gateway est le seul point d'entrée public de la plateforme (port 8008). "
        "Il centralise la sécurité, le routage et la gestion des erreurs. "
        "Construit avec FastAPI, il délègue la validation JWT à ms-auth-core."
    ))

    add_heading(doc, "6.1 Mécanisme d'Authentification", level=2)
    add_para(doc, "Pour chaque requête protégée :")
    add_bullet(doc, "1. Extraction du header Authorization: Bearer <token>")
    add_bullet(doc, "2. Appel GET {AUTH_CORE_BASE_URL}/auth/me avec le token")
    add_bullet(doc, "3. Récupération des claims JWT (sub, roles, username)")
    add_bullet(doc, "4. Vérification du rôle via require_roles('role1', 'role2') (FastAPI Depends)")
    add_bullet(doc, "5. Si autorisé : proxification vers le service cible avec Authorization préservé")

    add_heading(doc, "6.2 Résilience", level=2)
    add_bullet(doc, "Retry automatique (2 tentatives) sur les requêtes GET avec backoff exponentiel (150ms, 300ms)")
    add_bullet(doc, "Timeout configurable par service (défaut 8s)")
    add_bullet(doc, "Propagation correcte des erreurs HTTP (400, 401, 403, 404, 422, 502, 503)")

    add_heading(doc, "6.3 CORS", level=2)
    add_para(doc, "Origines autorisées : http://localhost:5173 (Vite dev), http://localhost:3000 (prod Docker), http://localhost:4173 (Vite preview). Headers exposés : x-request-id, x-correlation-id.")

    add_heading(doc, "6.4 Routes Exposées", level=2)
    routes = [
        ("GET", "/gateway/health", "Public", "Healthcheck du gateway"),
        ("GET", "/api/me", "Authentifié", "Retourne les claims JWT de l'utilisateur courant"),
        ("GET/POST/PUT/DELETE", "/api/content/courses[/*]", "teacher, admin", "CRUD cours et assets"),
        ("GET/POST", "/api/access/courses[/*]", "student", "Accès cours + presigned URLs"),
        ("GET/PATCH/POST", "/api/admin/users[/*]", "admin", "Gestion utilisateurs"),
        ("GET/PATCH", "/api/notifications[/*]", "Authentifié", "Notifications utilisateur"),
        ("GET/POST/PATCH", "/api/stats", "Authentifié", "Statistiques par rôle"),
        ("POST/GET/PATCH/DELETE", "/api/calendar/events[/*]", "Selon rôle", "Calendrier académique"),
        ("POST/GET", "/api/forum/threads[/*]", "Authentifié", "Forum et messages"),
        ("POST/GET", "/api/assignments[/*]", "Selon rôle", "Devoirs, soumissions, notes"),
        ("POST/GET", "/api/ai/*", "Authentifié", "Assistant IA"),
    ]
    rt = doc.add_table(rows=1 + len(routes), cols=4)
    rt.style = 'Table Grid'
    headers = ['Méthode', 'Chemin', 'Rôle Requis', 'Description']
    for i, h in enumerate(headers):
        set_cell_bg(rt.rows[0].cells[i], '1A56DB')
        p = rt.rows[0].cells[i].paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.color.rgb = BLANC; r.font.size = Pt(9.5)
    for ri, (method, path, role, desc) in enumerate(routes):
        bg = 'EFF6FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate([method, path, role, desc]):
            set_cell_bg(rt.rows[ri+1].cells[ci], bg)
            p = rt.rows[ri+1].cells[ci].paragraphs[0]
            r = p.add_run(val); r.font.size = Pt(9)
            r.font.color.rgb = BLEU_MED if ci == 0 else GRIS_TEXTE
            if ci == 0: r.bold = True
    doc.add_paragraph()

    add_page_break(doc)

    # ── 7. FRONTEND ───────────────────────────────────────────────────────────
    add_heading(doc, "7. Interface Utilisateur (Frontend)", level=1)
    add_para(doc, (
        "Le frontend est une Single Page Application (SPA) développée avec React 18, TypeScript et Vite. "
        "Le design system shadcn/ui sur Tailwind CSS assure une interface moderne, accessible et cohérente. "
        "L'application est conteneurisée via un build multi-stage (Vite → nginx) et servie sur le port 3000."
    ))

    add_heading(doc, "7.1 Stack Technique Frontend", level=2)
    stack = [
        ("React 18 + TypeScript", "Composants typés, hooks, context API"),
        ("Vite", "Build ultra-rapide, HMR en développement"),
        ("Tailwind CSS", "Utility-first CSS, design responsive"),
        ("shadcn/ui", "Composants UI accessibles (Button, Card, Dialog, Table, Tabs, Badge, etc.)"),
        ("React Router v6", "Navigation SPA, routes protégées par rôle"),
        ("nginx", "Serveur statique en production, proxy /api vers le gateway"),
    ]
    st = doc.add_table(rows=1 + len(stack), cols=2)
    st.style = 'Table Grid'
    for i, h in enumerate(['Technologie', 'Rôle']):
        set_cell_bg(st.rows[0].cells[i], '0F294D')
        p = st.rows[0].cells[i].paragraphs[0]
        r = p.add_run(h); r.bold = True; r.font.color.rgb = BLANC; r.font.size = Pt(10)
    for ri, (tech, role) in enumerate(stack):
        bg = 'E8F0FE' if ri % 2 == 0 else 'FFFFFF'
        set_cell_bg(st.rows[ri+1].cells[0], bg); set_cell_bg(st.rows[ri+1].cells[1], bg)
        r0 = st.rows[ri+1].cells[0].paragraphs[0].add_run(tech)
        r0.font.size = Pt(9.5); r0.bold = True; r0.font.color.rgb = BLEU_FONCE
        r1 = st.rows[ri+1].cells[1].paragraphs[0].add_run(role)
        r1.font.size = Pt(9.5); r1.font.color.rgb = GRIS_TEXTE
    doc.add_paragraph()

    add_heading(doc, "7.2 Pages de l'Application", level=2)
    pages = [
        ("LoginPage", "Toutes", "Authentification Keycloak OIDC, redirect post-login"),
        ("DashboardPage", "Toutes", "Tableau de bord rôle-adaptatif, stats en temps-réel, fil d'activité forum"),
        ("AdminPage", "admin", "Gestion utilisateurs : liste, création, modification des rôles"),
        ("TeacherPage", "teacher, admin", "Gestion des cours, upload d'assets, publication de devoirs"),
        ("StudentPage", "student", "Catalogue cours, téléchargement assets, rendu de devoirs"),
        ("CalendarPage", "Toutes", "Calendrier mensuel, filtrage par module, création (teacher/admin)"),
        ("ForumPage", "Toutes", "Fils de discussion, chat WebSocket temps-réel par salle"),
        ("ExamPage", "Toutes", "Devoirs publiés, soumission fichiers/texte, notes (teacher)"),
        ("NotificationsPage", "Toutes", "Centre de notifications, marquage lu/tout lu"),
        ("ProfilePage", "Toutes", "Profil utilisateur, déconnexion"),
        ("AssistantPage", "Toutes", "Interface chat IA, résumé, FAQ"),
    ]
    pt = doc.add_table(rows=1 + len(pages), cols=3)
    pt.style = 'Table Grid'
    for i, h in enumerate(['Page', 'Rôles', 'Fonctionnalité']):
        set_cell_bg(pt.rows[0].cells[i], '1A56DB')
        r = pt.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = BLANC; r.font.size = Pt(10)
    for ri, (page, roles, fn) in enumerate(pages):
        bg = 'EFF6FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate([page, roles, fn]):
            set_cell_bg(pt.rows[ri+1].cells[ci], bg)
            r = pt.rows[ri+1].cells[ci].paragraphs[0].add_run(val)
            r.font.size = Pt(9.5); r.font.color.rgb = BLEU_FONCE if ci == 0 else GRIS_TEXTE
            if ci == 0: r.bold = True
    doc.add_paragraph()

    add_heading(doc, "7.3 Authentification Frontend", level=2)
    add_para(doc, (
        "L'auth-context expose un hook useAuth() qui gère le token JWT, "
        "le refresh automatique et la fonction apiFetch() qui ajoute automatiquement "
        "le header Authorization: Bearer <token> à chaque appel API. "
        "ProtectedRoute redirige vers /login si non authentifié. RoleRoute affiche /unauthorized si le rôle est insuffisant."
    ))

    add_page_break(doc)

    # ── 8. SÉCURITÉ ───────────────────────────────────────────────────────────
    add_heading(doc, "8. Sécurité et Authentification", level=1)

    add_heading(doc, "8.1 Modèle de Sécurité", level=2)
    add_para(doc, (
        "La sécurité repose sur le standard OAuth2/OpenID Connect avec Keycloak comme serveur d'autorisation. "
        "Les JWT sont signés avec RS256 (clé asymétrique). Aucun secret partagé entre services."
    ))

    add_heading(doc, "8.2 RBAC — Contrôle d'Accès Basé sur les Rôles", level=2)
    rbac = [
        ("admin", "Gestion complète : utilisateurs, cours, calendrier, toutes les pages"),
        ("teacher", "Création cours/assets, publication devoirs, gestion calendrier, forum"),
        ("student", "Lecture cours, téléchargement assets, soumission devoirs, forum, calendrier"),
    ]
    for role, perms in rbac:
        p = doc.add_paragraph()
        r1 = p.add_run(f"  {role.upper()} : ")
        r1.bold = True; r1.font.color.rgb = BLEU_MED; r1.font.size = Pt(10.5)
        r2 = p.add_run(perms)
        r2.font.color.rgb = GRIS_TEXTE; r2.font.size = Pt(10.5)

    add_heading(doc, "8.3 Flux OAuth2/OIDC", level=2)
    add_bullet(doc, "1. L'utilisateur se connecte via la LoginPage (Authorization Code Flow)")
    add_bullet(doc, "2. Keycloak émet un id_token + access_token (JWT)")
    add_bullet(doc, "3. Le frontend stocke le token et l'envoie dans chaque requête API")
    add_bullet(doc, "4. Le gateway valide le token auprès de ms-auth-core (/auth/me)")
    add_bullet(doc, "5. ms-auth-core vérifie la signature via JWKS et retourne les claims")
    add_bullet(doc, "6. Le gateway inspecte realm_access.roles et autorise/refuse selon la route")

    add_heading(doc, "8.4 Sécurité des Communications", level=2)
    add_bullet(doc, "Réseau Docker interne ent-network : les services ne sont pas exposés directement")
    add_bullet(doc, "CORS strict : seules les origines connues sont autorisées")
    add_bullet(doc, "Token interne (INTERNAL_API_TOKEN) pour les appels inter-services sans JWT utilisateur")
    add_bullet(doc, "Mots de passe hachés par Keycloak (bcrypt)")
    add_bullet(doc, "Variables sensibles isolées dans .env.compose (non versionné)")

    add_page_break(doc)

    # ── 9. OBSERVABILITÉ ──────────────────────────────────────────────────────
    add_heading(doc, "9. Observabilité — Prometheus & Grafana", level=1)
    add_para(doc, (
        "Chaque microservice expose un endpoint /metrics (format Prometheus) via le module runtime.py partagé. "
        "Prometheus scrape tous les 9 services toutes les 15 secondes. Grafana visualise ces métriques "
        "avec la datasource Prometheus auto-provisionnée."
    ))

    add_heading(doc, "9.1 Métriques Collectées (runtime.py)", level=2)
    add_bullet(doc, "Compteur de requêtes HTTP par méthode, route et code de statut")
    add_bullet(doc, "Histogramme de latence des requêtes (percentiles p50, p95, p99)")
    add_bullet(doc, "Compteur d'erreurs par type d'exception")
    add_bullet(doc, "Métriques système (CPU, mémoire) via prometheus-client")

    add_heading(doc, "9.2 Configuration Prometheus", level=2)
    add_para(doc, "Services scrappés (scrape_interval: 15s) :")
    services_prom = [
        "ms-auth-core:8000/metrics",
        "ms-api-gateway:8000/metrics",
        "ms-identity-admin:8000/metrics",
        "ms-course-content:8000/metrics",
        "ms-course-access:8000/metrics",
        "ms-notification:8000/metrics",
        "ms-calendar-schedule:8000/metrics",
        "ms-forum-chat:8000/metrics",
        "ms-exam-assignment:8000/metrics",
    ]
    for s in services_prom:
        add_bullet(doc, s)

    add_heading(doc, "9.3 Grafana", level=2)
    add_para(doc, (
        "Grafana (port 3001) est pré-configuré avec la datasource Prometheus via le fichier "
        "infra/compose/grafana/provisioning/datasources/datasource.yml. Un dashboard ENT Platform "
        "est provisionné automatiquement depuis infra/compose/grafana/provisioning/dashboards/ent-platform.json."
    ))

    add_page_break(doc)

    # ── 10. TESTS ─────────────────────────────────────────────────────────────
    add_heading(doc, "10. Tests et Validation", level=1)
    add_para(doc, (
        "La plateforme inclut des scripts de smoke-test bash dans le dossier scripts/ "
        "qui valident l'ensemble du pipeline fonctionnel en conditions réelles (services démarrés)."
    ))

    scripts = [
        ("compose-healthcheck.sh", "Vérifie que tous les conteneurs Docker sont healthy"),
        ("rbac-smoke-test.sh", "Teste les permissions par rôle (admin/teacher/student) sur le gateway"),
        ("course-content-smoke-test.sh", "CRUD complet cours + upload asset (teacher)"),
        ("course-access-smoke-test.sh", "Listing cours + génération presigned URL (student)"),
        ("identity-admin-smoke-test.sh", "Création utilisateur + modification rôle (admin)"),
        ("calendar-smoke-test.sh", "CRUD événements calendrier + vérification RabbitMQ"),
        ("forum-smoke-test.sh", "Création fil + messages + connexion WebSocket"),
        ("exam-smoke-test.sh", "Création devoir + soumission + notation + notification"),
    ]
    st = doc.add_table(rows=1 + len(scripts), cols=2)
    st.style = 'Table Grid'
    for i, h in enumerate(['Script', 'Périmètre Testé']):
        set_cell_bg(st.rows[0].cells[i], '059669')
        r = st.rows[0].cells[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = BLANC; r.font.size = Pt(10)
    for ri, (script, desc) in enumerate(scripts):
        bg = 'F0FDF4' if ri % 2 == 0 else 'FFFFFF'
        set_cell_bg(st.rows[ri+1].cells[0], bg); set_cell_bg(st.rows[ri+1].cells[1], bg)
        r0 = st.rows[ri+1].cells[0].paragraphs[0].add_run(script)
        r0.font.name = 'Courier New'; r0.font.size = Pt(9); r0.font.color.rgb = VERT
        r1 = st.rows[ri+1].cells[1].paragraphs[0].add_run(desc)
        r1.font.size = Pt(9.5); r1.font.color.rgb = GRIS_TEXTE
    doc.add_paragraph()

    add_page_break(doc)

    # ── 11. DÉPLOIEMENT ───────────────────────────────────────────────────────
    add_heading(doc, "11. Déploiement avec Docker Compose", level=1)
    add_para(doc, (
        "L'ensemble de la plateforme est défini dans un fichier docker-compose.yml unique. "
        "La commande de démarrage standard est docker compose --env-file .env.compose up -d --build."
    ))

    add_heading(doc, "11.1 Ordre de Démarrage et Dépendances", level=2)
    add_para(doc, "Docker Compose gère automatiquement l'ordre grâce aux depends_on avec condition:")
    add_bullet(doc, "1. postgres-keycloak (healthy) → keycloak (started)")
    add_bullet(doc, "2. keycloak → keycloak-realm-bootstrap (completed) → keycloak-users-bootstrap (completed)")
    add_bullet(doc, "3. cassandra (healthy) → cassandra-init (completed)")
    add_bullet(doc, "4. minio (started) → minio-init (completed)")
    add_bullet(doc, "5. rabbitmq (healthy)")
    add_bullet(doc, "6. ms-auth-core (healthy) → tous les microservices")
    add_bullet(doc, "7. Tous les microservices (healthy) → ms-api-gateway (healthy)")
    add_bullet(doc, "8. ms-api-gateway (healthy) → frontend")

    add_heading(doc, "11.2 Volumes Persistants", level=2)
    vols = ["keycloak_postgres_data", "cassandra_data", "rabbitmq_data", "minio_data", "redis_data", "grafana_data", "ollama_data"]
    for v in vols:
        add_bullet(doc, v)

    add_heading(doc, "11.3 Structure des Dockerfiles", level=2)
    add_para(doc, "Pattern Python slim (inspiré de ms-notification/Dockerfile) :")
    add_bullet(doc, "FROM python:3.12-slim")
    add_bullet(doc, "WORKDIR /app, COPY requirements.txt, RUN pip install --no-cache-dir")
    add_bullet(doc, "COPY app/ ./app/")
    add_bullet(doc, "CMD [\"uvicorn\", \"app.main:app\", \"--host\", \"0.0.0.0\", \"--port\", \"8000\"]")

    add_heading(doc, "11.4 Frontend Multi-Stage Build", level=2)
    add_bullet(doc, "Stage 1 (builder) : node:20-alpine, npm ci, npm run build → /app/dist")
    add_bullet(doc, "Stage 2 (production) : nginx:alpine, copie /app/dist, nginx.conf (proxy /api)")

    add_page_break(doc)

    # ── 12. PATTERNS ──────────────────────────────────────────────────────────
    add_heading(doc, "12. Patterns et Conventions de Développement", level=1)

    add_heading(doc, "12.1 runtime.py — Module Partagé", level=2)
    add_para(doc, (
        "Chaque service importe setup_service_runtime(app, service_name) depuis son runtime.py local "
        "(copie identique pour chaque service). Ce module configure :"
    ))
    add_bullet(doc, "Logging structuré JSON (corrélation request_id par requête)")
    add_bullet(doc, "Middleware Prometheus : compteur + histogramme automatiques")
    add_bullet(doc, "Gestionnaire d'exceptions global (retourne JSON avec detail)")
    add_bullet(doc, "Endpoint GET /metrics (exposition format Prometheus)")

    add_heading(doc, "12.2 Pattern Cassandra _ensure_db()", level=2)
    add_para(doc, "Initialisation paresseuse : appelée une seule fois au premier appel API, crée le keyspace et les tables IF NOT EXISTS. Évite les erreurs au démarrage si Cassandra n'est pas encore prête.")

    add_heading(doc, "12.3 Pattern d'Authentification Inter-Services", level=2)
    add_para(doc, "Les services qui appellent un autre service (ex: ms-course-access → ms-course-content) utilisent soit le JWT utilisateur propagé soit un INTERNAL_API_TOKEN pour les appels système.")

    add_heading(doc, "12.4 Versionnement des Événements", level=2)
    add_para(doc, "Tous les types d'événements incluent une version (ex: user.created.v1). La clé de routage RabbitMQ = {domain}.{action} (ex: user.created). Cela permet l'évolution indépendante des consommateurs.")

    add_page_break(doc)

    # ── 13. CONCLUSION ────────────────────────────────────────────────────────
    add_heading(doc, "13. Conclusion et Perspectives", level=1)

    add_heading(doc, "13.1 Bilan du Projet", level=2)
    add_para(doc, (
        "Le projet ENT-EST-Salé constitue une implémentation complète et fonctionnelle d'une plateforme "
        "numérique académique selon les principes DevOps et Cloud. L'ensemble des 10 microservices, "
        "l'infrastructure à 9 composants, le frontend multi-pages et la stack d'observabilité ont été "
        "réalisés et intégrés dans un environnement Docker Compose entièrement automatisé."
    ))

    realisations = [
        "Architecture microservices avec découplage fort et interfaces API claires",
        "Sécurité industrielle : OAuth2/OIDC, RBAC, JWT RS256, communications chiffrées",
        "Persistance NoSQL scalable avec Apache Cassandra",
        "Messagerie événementielle asynchrone avec RabbitMQ",
        "Stockage objet cloud-native avec MinIO (S3-compatible)",
        "Interface utilisateur moderne et responsive (React + shadcn/ui)",
        "Chat temps-réel via WebSocket",
        "Observabilité complète avec Prometheus et Grafana",
        "Scripts de smoke-test couvrant tous les flux fonctionnels",
        "Déploiement one-command via Docker Compose",
    ]
    for r in realisations:
        add_bullet(doc, f"✓  {r}")

    add_heading(doc, "13.2 Perspectives d'Évolution", level=2)
    perspectives = [
        "Migration vers Kubernetes (Helm charts déjà préparés dans infra/manifests/)",
        "Ajout de dashboards Grafana avancés (latence p99, taux d'erreur par service)",
        "Implémentation du service ms-ai-assistant en production (GPU cloud)",
        "CI/CD pipeline GitLab/GitHub Actions pour le build et les tests automatiques",
        "Ajout d'un service de gestion des emplois du temps",
        "Application mobile React Native connectée au même backend",
        "Haute disponibilité : Cassandra multi-nœuds, RabbitMQ cluster, MinIO distributed",
    ]
    for p in perspectives:
        add_bullet(doc, p)

    add_para(doc, "")
    add_para(doc, (
        "Ce projet démontre la maîtrise de l'écosystème DevOps et Cloud dans un contexte académique réel, "
        "en appliquant les bonnes pratiques de l'industrie : Infrastructure as Code, Containers-as-a-Service, "
        "12-Factor App, observabilité native et sécurité by design."
    ), bold=False)

    return doc

# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    doc = build_document()
    out = "/home/khalidmarzoug/Documents/devops-project/ent-est-sale-platform/Rapport_ENT_EST_Sale.docx"
    doc.save(out)
    print(f"✅  Rapport sauvegardé : {out}")
